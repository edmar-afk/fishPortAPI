from docx import Document
from django.http import HttpResponse

def generate_fishing_permit_docx(permit):
    # Load the template
    template_path = "./FISHING-PERMIT.docx"  # Update this path
    doc = Document(template_path)

    # Replace placeholders in the document
    placeholders = {
        "{owner_name}": permit.owner_name,
        "{address}": permit.address,
        "{homeport}": permit.home_port,
        "{vessel_name}": permit.vessel_name,
        "{vessel_type}": permit.vessel_type,
        "{color}": permit.color,
        "{service_type}": permit.service_type,
        "{vessel_description}": permit.vessel_description,
        "{length}": permit.length,
        "{breadth}": permit.breadth,
        "{depth}": permit.depth,
        "{gross}": permit.gross,
        "{net}": permit.net,
        "{engine}": permit.engine,
        "{serial_num}": permit.serial_num,
        "{horse_power}": permit.horse_power,
        "{cylinder_num}": permit.cylinder_num,
        "{engine_num}": permit.engine_num,
        "{crew_num}": permit.crew_num,
        "{coast_guard_num}": permit.coast_guard_num,
        "{mfvr_num}": permit.mfvr_num,
        "{or_num}": permit.or_num,
        "{date_issued}": permit.date_issued.strftime("%Y-%m-%d") if permit.date_issued else "",
        "{amount}": permit.amount,
        "{fishing_gear_used}": permit.fishing_gear_used,
    }

    # Replace text in the Word document
    for paragraph in doc.paragraphs:
        for placeholder, value in placeholders.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)

    # Handle tables (if placeholders exist in tables)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, value in placeholders.items():
                    if placeholder in cell.text:
                        cell.text = cell.text.replace(placeholder, value)

    # Generate response
    response = HttpResponse(content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    response["Content-Disposition"] = f"attachment; filename=FishingPermit-{permit.id}.docx"
    doc.save(response)
    return response
