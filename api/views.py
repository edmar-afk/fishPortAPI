# views.py
from rest_framework import generics, permissions, views
from rest_framework.permissions import AllowAny, IsAuthenticated
from django.contrib.auth.models import User
from rest_framework import status
from .serializers import UserSerializer, FishTypeSerializer, WeighInSerializer, ExpirationDateSerializer, WeighInHistorySerializer, VesselRegistrationSerializer, FishingPermitSerializer
from rest_framework.exceptions import ValidationError
from rest_framework.response import Response
from rest_framework.views import APIView
from django.http import FileResponse, Http404
from django.shortcuts import get_object_or_404
from django.views import View
from rest_framework.decorators import api_view
from .models import FishType, WeighIn, FishingPermit, VesselRegistration, ExpirationDate
from django.db.models import Sum
from django.utils import timezone
from datetime import datetime
from django.utils.timezone import now
from django.contrib.auth.hashers import make_password
import logging
from rest_framework.decorators import permission_classes
logger = logging.getLogger(__name__)
from datetime import timedelta
from django.db.models import Count  # Import Count for aggregation
from rest_framework.exceptions import NotFound
from rest_framework.decorators import action
from rest_framework import viewsets
from rest_framework.exceptions import ValidationError, NotFound  # Import exceptions
import os
from django.conf import settings
from docx import Document
from django.http import HttpResponse
from io import BytesIO
from django.http import JsonResponse

def generate_fishing_permit_docx(permit):
    # Debugging to check the state of the permit
    print("Generating DOCX for permit:", permit)

    # Build the path to the template file
    template_path = os.path.join(settings.BASE_DIR, 'FISHING-PERMIT.docx')

    # Ensure the file exists
    if not os.path.exists(template_path):
        raise FileNotFoundError(f"Template not found: {template_path}")

    # Load the template document
    doc = Document(template_path)

    # Handle paragraphs, replacing placeholders in each run
    for paragraph in doc.paragraphs:
        # Loop through each run in the paragraph
        for run in paragraph.runs:
            # Check if the run contains a placeholder
            if '{' in run.text and '}' in run.text:
                print(f"Replacing placeholders in run: {run.text}")
                
                # Print values of the fields for debugging
                print(f"Service type: {permit.service_type}")
                print(f"Coast Guard Number: {permit.coast_guard_num}")
                print(f"MFVR Number: {permit.mfvr_num}")
                print(f"OR Number: {permit.or_num}")
                print(f"Date Issued: {permit.date_issued}")
                print(f"Fishing Gear Used: {permit.fishing_gear_used}")

                run.text = run.text.format(
                    owner_name=permit.owner_name,
                    address=permit.address,
                    homeport=permit.home_port,
                    vessel_name=permit.vessel_name,
                    vessel_type=permit.vessel_type,
                    color=permit.color,       
                    vessel_description=permit.vessel_description,
                    length=permit.length,
                    breadth=permit.breadth,
                    depth=permit.depth,
                    gross=permit.gross,
                    net=permit.net,
                    engine=permit.engine,
                    serial_num=permit.serial_num,
                    horse_power=permit.horse_power,
                    cylinder_num=permit.cylinder_num,
                    engine_num=permit.engine_num,
                    crew_num=permit.crew_num,
                    amount=permit.amount,
                     
                    # below here is not displaying
                    service_type=permit.service_type,
                    coast_guard_num=permit.coast_guard_num,
                    mfvr_num=permit.mfvr_num,
                    or_num=permit.or_num,
                    date_issued=permit.date_issued.strftime('%Y-%m-%d') if permit.date_issued else '',
                    fishing_gear_used=permit.fishing_gear_used,
                )

    # Handle tables in the document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        # Debugging: print the text of each run in the table cell
                        print(f"Checking table cell run: {run.text}")
                        
                        # Replace placeholders in table cells
                        if '{' in run.text and '}' in run.text:
                            print(f"Replacing placeholders in table cell: {run.text}")
                            
                            # Print values of the fields for debugging
                            print(f"Service type: {permit.service_type}")
                            print(f"Coast Guard Number: {permit.coast_guard_num}")
                            print(f"MFVR Number: {permit.mfvr_num}")
                            print(f"OR Number: {permit.or_num}")
                            print(f"Date Issued: {permit.date_issued}")
                            print(f"Fishing Gear Used: {permit.fishing_gear_used}")

                            run.text = run.text.format(
                                owner_name=permit.owner_name,
                                address=permit.address,
                                homeport=permit.home_port,
                                vessel_name=permit.vessel_name,
                                vessel_type=permit.vessel_type,
                                color=permit.color,
                                vessel_description=permit.vessel_description,
                                length=permit.length,
                                breadth=permit.breadth,
                                depth=permit.depth,
                                gross=permit.gross,
                                net=permit.net,
                                engine=permit.engine,
                                serial_num=permit.serial_num,
                                horse_power=permit.horse_power,
                                cylinder_num=permit.cylinder_num,
                                engine_num=permit.engine_num,
                                crew_num=permit.crew_num,
                                amount=permit.amount,
                                
                                
                                # below here is not displaying
                                service_type=permit.service_type,
                                coast_guard_num=permit.coast_guard_num,
                                mfvr_num=permit.mfvr_num,
                                or_num=permit.or_num,
                                date_issued=permit.date_issued.strftime('%Y-%m-%d') if permit.date_issued else 'none',
                                fishing_gear_used=permit.fishing_gear_used,
                            )

    # Save to an in-memory file
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)  # Rewind the buffer to the start so it can be read

    # Create an HTTP response to download the file
    response = HttpResponse(
        buffer,
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename=Fishing-Permit-{permit.owner_name}.docx'

    return response

def download_fishing_permit(request, permit_id):
    # Fetch the fishing permit object
    permit = get_object_or_404(FishingPermit, id=permit_id)

    # Generate and return the DOCX response
    return generate_fishing_permit_docx(permit)









@api_view(['POST'])
@permission_classes([AllowAny])
def register_user(request):
    if request.method == 'POST':
        # Print the incoming request data to the console
        print("Received data:", request.data)

        # Continue with your serializer logic
        serializer = UserSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response(serializer.data, status=status.HTTP_201_CREATED)

        # If there are errors, print them too
        print("Validation errors:", serializer.errors)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)



class FishTypeListCreateView(generics.ListCreateAPIView):
    queryset = FishType.objects.all()
    serializer_class = FishTypeSerializer
    permission_classes = [AllowAny]


# New view for listing FishTypes only
class FishTypeListView(generics.ListAPIView):
    queryset = FishType.objects.all()
    serializer_class = FishTypeSerializer
    permission_classes = [AllowAny]

# New view for listing FishTypes only
@api_view(['DELETE'])
@permission_classes([AllowAny])
def delete_fish_type(request, pk):
    try:
        fish_type = FishType.objects.get(pk=pk)
    except FishType.DoesNotExist:
        return Response({'error': 'Fish type not found.'}, status=status.HTTP_404_NOT_FOUND)

    fish_type.delete()
    return Response({'message': 'Fish type deleted successfully.'}, status=status.HTTP_204_NO_CONTENT)

class WeighInCreateView(generics.CreateAPIView):
    queryset = WeighIn.objects.all()
    serializer_class = WeighInSerializer
    permission_classes = [AllowAny]

class WeighInListAPIView(generics.ListAPIView):
    queryset = WeighIn.objects.all()
    serializer_class = WeighInHistorySerializer
    permission_classes = [AllowAny]

class FishTotalKilosView(APIView):
    permission_classes = [AllowAny]
    def get(self, request, *args, **kwargs):
        # Get the current date
        today = timezone.now().date()

        # Filter weigh-ins for the current date and aggregate total kg per fish type
        fish_totals = (
            WeighIn.objects
            .filter(date_weighin__date=today)
            .values('fish')
            .annotate(total_kg=Sum('kg'))
            .order_by('fish')
        )

        # Prepare the response data
        response_data = []
        for item in fish_totals:
            fish_id = item['fish']
            total_kg = item['total_kg']
            fish_type = FishType.objects.get(id=fish_id)
            response_data.append({
                'fish': fish_type.name,
                'total_kg': total_kg
            })

        return Response(response_data, status=status.HTTP_200_OK)

class UserDetailView(generics.RetrieveAPIView):
    queryset = User.objects.all()
    serializer_class = UserSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_object(self):
        return self.request.user




class NonSuperUserCountAPIView(APIView):
    permission_classes = [AllowAny]
    def get(self, request, *args, **kwargs):
        # Count users who are not superusers
        non_superuser_count = User.objects.filter(is_superuser=False).count()
        return Response({'non_superuser_count': non_superuser_count})


class FishermanListView(generics.ListAPIView):
    permission_classes = [AllowAny]

    queryset = User.objects.filter(is_superuser=False)
    serializer_class = UserSerializer

    def get_queryset(self):
        # Optionally, you can customize this method to filter the users
        return User.objects.filter(is_superuser=False)


class FishingPermitCreateAPIView(APIView):
    permission_classes = [AllowAny]

    def post(self, request, user_id):
        # Retrieve the User object based on the provided user_id
        try:
            owner = User.objects.get(id=user_id)
        except User.DoesNotExist:
            return Response({"error": "User not found"}, status=status.HTTP_404_NOT_FOUND)

        # Include the owner object and set owner_name to user.first_name in the form data
        data = request.data.copy()
        data['owner'] = owner.id  # Set the owner to the user's ID
        data['owner_name'] = owner.first_name  # Automatically set owner_name to first_name

        # Print the submitted data for debugging
        print("Data being submitted:", data)

        # Serialize the data
        serializer = FishingPermitSerializer(data=data)
        if serializer.is_valid():
            serializer.save()  # Save the FishingPermit instance
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            # Print and return validation errors for debugging
            print("Validation errors:", serializer.errors)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)




class LatestFishingPermitAPIView(generics.RetrieveAPIView):
    serializer_class = FishingPermitSerializer
    permission_classes = [AllowAny]

    def get_object(self):
        user_id = self.kwargs['userId']
        # Fetch the latest permit for the user if it exists
        return FishingPermit.objects.filter(owner_id=user_id).last()


class FishingPermitDetailsView(generics.RetrieveAPIView):
    permission_classes = [AllowAny]
    serializer_class = FishingPermitSerializer

    def get_object(self):
        user_id = self.kwargs.get("userId")
        # Get the latest fishing permit for the user, ordered by ID in descending order
        permit = FishingPermit.objects.filter(owner__id=user_id).order_by('-id').first()
        if permit is None:
            self.raise_exception()  # Raise a 404 error if no permit is found
        return permit



class VesselRegistrationCreateAPIView(APIView):
    permission_classes = [AllowAny]

    def post(self, request, user_id):
        # Retrieve the User object based on the provided user_id
        try:
            owner = User.objects.get(id=user_id)
        except User.DoesNotExist:
            return Response({"error": "User not found"}, status=status.HTTP_404_NOT_FOUND)

        # Include the owner object in the form data
        data = request.data.copy()
        data['owner'] = owner.id  # Set the owner to the user's ID

        # Print the submitted data for debugging
        print("Data being submitted:", data)

        # Serialize the data
        serializer = VesselRegistrationSerializer(data=data)
        if serializer.is_valid():
            serializer.save()  # Save the VesselRegistration instance
            return Response(serializer.data, status=status.HTTP_201_CREATED)
        else:
            # Print and return validation errors for debugging
            print("Validation errors:", serializer.errors)
            return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)


class LatestVesselRegAPIView(generics.RetrieveAPIView):
    serializer_class = VesselRegistrationSerializer
    permission_classes = [AllowAny]

    def get_object(self):
        user_id = self.kwargs['userId']
        # Fetch the latest permit for the user if it exists
        return VesselRegistration.objects.filter(owner_id=user_id).last()

class VesselRegistrationDetailsView(generics.RetrieveAPIView):
    permission_classes = [AllowAny]
    serializer_class = VesselRegistrationSerializer

    def get_object(self):
        vessel_id = self.kwargs.get("vesselId")  # Extract `vesselId` from URL
        vessel = VesselRegistration.objects.filter(id=vessel_id).first()
        if vessel is None:
            raise NotFound("Vessel not found.")  # Use NotFound exception for 404
        return vessel



class TotalWeightTodayView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        today = now().date()  # Get the current date

        # Query WeighIn for today only and aggregate the total kg
        total_weight = WeighIn.objects.filter(
            date_weighin__date=today
        ).aggregate(total_weight=Sum('kg'))['total_weight'] or 0

        return Response({"total_weight_today": total_weight})


class TotalPriceTodayView(APIView):
    permission_classes = [AllowAny]

    def get(self, request):
        today = now().date()  # Get the current date

        # Query WeighIn for today only and aggregate the total price
        total_price = WeighIn.objects.filter(
            date_weighin__date=today
        ).aggregate(total_price=Sum('total_price'))['total_price'] or 0

        return Response({"total_price_today": total_price})



class WeighInByFishView(generics.GenericAPIView):
    permission_classes = [AllowAny]

    def get(self, request, *args, **kwargs):
        fish_name = self.kwargs['fish_name']
        current_date = timezone.now().date()

        try:
            fish = FishType.objects.get(name=fish_name)
        except FishType.DoesNotExist:
            raise NotFound(detail="Fish type not found.")

        weighin_queryset = WeighIn.objects.filter(fish=fish, date_weighin__date=current_date)
        total_kg = weighin_queryset.aggregate(total_kg=Sum('kg'))['total_kg'] or 0

        return Response({
            "fish_name": fish_name,
            "total_kg_today": total_kg
        })



class UserDeleteView(APIView):
    permission_classes = [AllowAny]
    """
    API view to delete a user based on their ID.
    """
    def delete(self, request, *args, **kwargs):
        # Retrieve the user ID from the URL
        user_id = kwargs.get('id')

        # Try to get the user object
        try:
            user = User.objects.get(id=user_id)
        except User.DoesNotExist:
            return Response(
                {"detail": "User not found."}, status=status.HTTP_404_NOT_FOUND
            )

        # Check if the user is allowed to delete (optional)
        # For example, prevent deleting superusers or the current authenticated user
        if user.is_superuser or user == request.user:
            return Response(
                {"detail": "You cannot delete this user."},
                status=status.HTTP_403_FORBIDDEN,
            )

        # Delete the user
        user.delete()
        return Response({"detail": "User deleted successfully."}, status=status.HTTP_204_NO_CONTENT)
    
    
    
    
class GrantFishingPermitView(APIView):
    permission_classes = [AllowAny]
    def patch(self, request, permitId):
        # Fetch the fishing permit object or return 404 if not found
        fishing_permit = get_object_or_404(FishingPermit, id=permitId)
        
        # Update the status field
        fishing_permit.status = 'Granted'
        fishing_permit.save()
        
        # Serialize the updated object
        serializer = FishingPermitSerializer(fishing_permit)
        
        # Return the serialized data
        return Response(serializer.data, status=status.HTTP_200_OK)
    

class GrantVesselRegView(APIView):
    permission_classes = [AllowAny]
    def patch(self, request, vesselId):
        # Fetch the fishing permit object or return 404 if not found
        vessel_registration = get_object_or_404(VesselRegistration, id=vesselId)
        
        # Update the status field
        vessel_registration.status = 'Granted'
        vessel_registration.save()
        
        # Serialize the updated object
        serializer = VesselRegistrationSerializer(vessel_registration)
        
        # Return the serialized data
        return Response(serializer.data, status=status.HTTP_200_OK)
    
    

class FishingPermitStatusView(APIView):
    permission_classes = [AllowAny]
    """
    API view to check the latest status of a fishing permit by ownerId.
    """
    def get(self, request, ownerId):
        # Validate that the owner exists
        owner = get_object_or_404(User, id=ownerId)

        # Fetch the latest permit for the given owner
        latest_permit = FishingPermit.objects.filter(owner=owner).order_by('-id').first()

        if latest_permit:
            return Response({
                "ownerId": ownerId,
                "owner_name": owner.get_full_name(),  # Assuming `get_full_name` exists
                "permit": {
                    "id": latest_permit.id,
                    "vessel_name": latest_permit.vessel_name,
                    "status": latest_permit.status,
                    "date_issued": latest_permit.date_issued,
                }
            }, status=status.HTTP_200_OK)
        else:
            return Response({
                "ownerId": ownerId,
                "owner_name": owner.get_full_name(),
                "message": "No permits found for this owner."
            }, status=status.HTTP_404_NOT_FOUND)

            

class VesselRegistrationStatusView(APIView):
    permission_classes = [AllowAny]
    """
    API view to check the latest status of vessel registrations by ownerId.
    """
    def get(self, request, ownerId):
        # Validate that the owner exists
        owner = get_object_or_404(User, id=ownerId)

        # Fetch the latest vessel registration for the given owner
        latest_vessel = VesselRegistration.objects.filter(owner=owner).order_by('-id').first()

        if latest_vessel:
            return Response({
                "ownerId": ownerId,
                "owner_name": owner.get_full_name(),  # Assuming `get_full_name` exists
                "vessel": {
                    "id": latest_vessel.id,
                    "builder_name": latest_vessel.builder_name,
                    "year_built": latest_vessel.year_built,
                    "status": latest_vessel.status,
                    "amount": latest_vessel.amount,
                }
            }, status=status.HTTP_200_OK)
        else:
            return Response({
                "ownerId": ownerId,
                "owner_name": owner.get_full_name(),
                "message": "No vessel registrations found for this owner."
            }, status=status.HTTP_404_NOT_FOUND)
            
            
class TotalAmountGrantedFishingPermits(APIView):
    permission_classes = [AllowAny]
    def get(self, request, *args, **kwargs):
        # Filter fishing permits with status 'Granted' and sum the 'amount' field
        total_amount = FishingPermit.objects.filter(status='Granted').aggregate(total_amount=Sum('amount'))['total_amount']
        
        # If no permits are found with the 'Granted' status, set total_amount to 0
        if total_amount is None:
            total_amount = 0

        # Return the total amount in the response
        return Response({
            'total_amount_granted': total_amount
        }, status=status.HTTP_200_OK)
        
class TotalAmountGrantedVesselRegistrations(APIView):
    permission_classes = [AllowAny]
    def get(self, request, *args, **kwargs):
        # Filter VesselRegistrations with status 'Active' and sum the 'amount' field
        total_amount = VesselRegistration.objects.filter(status='Granted').aggregate(total_amount=Sum('amount'))['total_amount']
        
        # If no registrations are found with the 'Active' status, set total_amount to 0
        if total_amount is None:
            total_amount = 0

        # Return the total amount in the response
        return Response({
            'total_amount_active': total_amount
        }, status=status.HTTP_200_OK)
        
        
class UserDeleteAPIView(APIView):
    permission_classes = [AllowAny]  # Ensuring the user is authenticated before deleting

    def delete(self, request, *args, **kwargs):
        user_id = kwargs.get('id')
        try:
            user = User.objects.get(id=user_id)
            user.delete()
            return Response({'detail': 'User deleted successfully'}, status=status.HTTP_204_NO_CONTENT)
        except User.DoesNotExist:
            return Response({'detail': 'User not found'}, status=status.HTTP_404_NOT_FOUND)
        
        
        

class VesselRegistrationListByOwner(APIView):
    permission_classes = [AllowAny]  # Add permission if needed
    
    def get(self, request, owner_id, format=None):
        # Filter VesselRegistration by owner_id
        registrations = VesselRegistration.objects.filter(owner_id=owner_id)
        
        # If no records found, return a 404 error
        if not registrations:
            return Response({"detail": "Not found."}, status=status.HTTP_404_NOT_FOUND)
        
        # Serialize the list of IDs
        registration_ids = registrations.values_list('id', flat=True)
        
        return Response({"registration_ids": list(registration_ids)}, status=status.HTTP_200_OK)
    

class UserProfileView(View):
    permission_classes = [AllowAny]
    def get(self, request, userId):
        try:
            user = User.objects.get(pk=userId)
            user_data = {
                "id": user.id,
                "username": user.username,
                "first_name": user.first_name,
                "is_superuser": user.is_superuser,
            }
            return JsonResponse(user_data)
        except User.DoesNotExist:
            return JsonResponse({"error": "User not found."}, status=404)
        
        
class UserProfileUpdateView(APIView):
    permission_classes = [AllowAny]
    """
    API View to update the user's first_name and username.
    """
    def put(self, request, userId):
        user = get_object_or_404(User, pk=userId)
        data = request.data
        
        # Validate and update fields
        if "username" in data:
            user.username = data["username"]
        if "first_name" in data:
            user.first_name = data["first_name"]
        
        user.save()
        return Response(
            {
                "message": "User profile updated successfully.",
                "user": {
                    "id": user.id,
                    "username": user.username,
                    "first_name": user.first_name,
                },
            },
            status=status.HTTP_200_OK
        )
        


class ExpirationDateUploadView(APIView):
    permission_classes = [AllowAny]
    def post(self, request, vesselid, *args, **kwargs):
        try:
            # Get the VesselRegistration object based on the vesselid
            vessel_reg = VesselRegistration.objects.get(id=vesselid)

            # Create the ExpirationDate object and automatically set date_expired 30 days after date_registered
            expiration_date = ExpirationDate(
                vessel_reg=vessel_reg,
                date_registered=timezone.now().date(),
                date_expired=timezone.now().date() + timedelta(days=30)
            )

            # Save the object
            expiration_date.save()

            # Serialize the ExpirationDate object
            serializer = ExpirationDateSerializer(expiration_date)

            return Response(serializer.data, status=status.HTTP_201_CREATED)
        
        except VesselRegistration.DoesNotExist:
            return Response({"detail": "Vessel not found."}, status=status.HTTP_404_NOT_FOUND)

    


class ExpirationDateByVesselIdView(APIView):
    permission_classes = [AllowAny]  # Ensure only authenticated users can access
    """
    API View to retrieve expiration dates based on vesselId.
    """
    def get(self, request, vessel_id):
        try:
            # Fetch the vessel registration based on the vessel_id
            vessel_reg = VesselRegistration.objects.get(id=vessel_id)
            
            # Get expiration dates associated with this vessel registration
            expiration_dates = ExpirationDate.objects.filter(vessel_reg=vessel_reg)

            # Serialize the expiration dates
            serializer = ExpirationDateSerializer(expiration_dates, many=True)
            
            return Response(serializer.data, status=status.HTTP_200_OK)
        
        except VesselRegistration.DoesNotExist:
            return Response(
                {"error": "Vessel registration not found."},
                status=status.HTTP_404_NOT_FOUND
            )
        except Exception as e:
            return Response(
                {"error": f"An unexpected error occurred: {str(e)}"},
                status=status.HTTP_500_INTERNAL_SERVER_ERROR
            )